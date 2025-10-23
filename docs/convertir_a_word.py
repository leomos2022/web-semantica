"""
Script para convertir el documento Markdown a formato Word
Ejecutar en Google Colab para generar archivo .docx
"""

# Instalar dependencias necesarias
!pip install python-docx markdown beautifulsoup4 python-markdown-math

import markdown
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import re
from bs4 import BeautifulSoup

def markdown_to_docx(markdown_file_path, output_path):
    """
    Convertir archivo Markdown a documento Word
    """
    # Leer archivo Markdown
    with open(markdown_file_path, 'r', encoding='utf-8') as f:
        markdown_content = f.read()
    
    # Convertir Markdown a HTML
    html = markdown.markdown(markdown_content, extensions=['extra', 'codehilite'])
    soup = BeautifulSoup(html, 'html.parser')
    
    # Crear documento Word
    doc = Document()
    
    # Configurar estilos
    setup_document_styles(doc)
    
    # Procesar contenido HTML
    process_html_content(doc, soup)
    
    # Guardar documento
    doc.save(output_path)
    print(f"✅ Documento guardado en: {output_path}")

def setup_document_styles(doc):
    """
    Configurar estilos personalizados para el documento
    """
    # Estilo para títulos principales
    title_style = doc.styles.add_style('Custom Title', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Arial'
    title_font.size = Pt(16)
    title_font.bold = True
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Estilo para subtítulos
    subtitle_style = doc.styles.add_style('Custom Subtitle', WD_STYLE_TYPE.PARAGRAPH)
    subtitle_font = subtitle_style.font
    subtitle_font.name = 'Arial'
    subtitle_font.size = Pt(14)
    subtitle_font.bold = True
    
    # Estilo para código
    code_style = doc.styles.add_style('Custom Code', WD_STYLE_TYPE.PARAGRAPH)
    code_font = code_style.font
    code_font.name = 'Courier New'
    code_font.size = Pt(10)

def process_html_content(doc, soup):
    """
    Procesar contenido HTML y agregarlo al documento Word
    """
    for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'pre', 'ul', 'ol']):
        if element.name == 'h1':
            p = doc.add_heading(element.get_text(), level=1)
        elif element.name == 'h2':
            p = doc.add_heading(element.get_text(), level=2)
        elif element.name == 'h3':
            p = doc.add_heading(element.get_text(), level=3)
        elif element.name == 'h4':
            p = doc.add_heading(element.get_text(), level=4)
        elif element.name == 'p':
            p = doc.add_paragraph(element.get_text())
        elif element.name == 'pre':
            p = doc.add_paragraph(element.get_text(), style='Custom Code')
        elif element.name in ['ul', 'ol']:
            for li in element.find_all('li'):
                p = doc.add_paragraph(li.get_text(), style='List Bullet')

# Instrucciones de uso
print("📋 INSTRUCCIONES PARA GENERAR DOCUMENTO WORD")
print("="*50)
print("1. Sube el archivo 'Documento_Foro_Backpropagation_WebSemantica.md' a Google Colab")
print("2. Ejecuta este script en una celda de código")
print("3. El archivo .docx se generará automáticamente")
print("4. Descarga el archivo desde la sección de archivos de Colab")
print("\n✅ Script listo para ejecutar en Google Colab")

# Ejemplo de uso (descomenta para ejecutar)
# markdown_to_docx('Documento_Foro_Backpropagation_WebSemantica.md', 'FinFlow_Intelligence_Foro.docx')